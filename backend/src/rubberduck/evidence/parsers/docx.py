"""DOCX parser using python-docx."""

import logging
from pathlib import Path

from rubberduck.evidence.parsers.base import BaseParser, ParseResult

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):
    def parse(self, file_path: Path, **kwargs) -> ParseResult:
        from docx import Document

        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)

        # Extract tables
        table_texts = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            table_texts.append("\n".join(rows))

        if table_texts:
            full_text += "\n\n--- TABLES ---\n\n" + "\n\n".join(table_texts)

        # Metadata
        props = doc.core_properties
        metadata = {
            "author": props.author,
            "created": str(props.created) if props.created else None,
            "modified": str(props.modified) if props.modified else None,
            "title": props.title,
            "subject": props.subject,
            "category": props.category,
            "revision": props.revision,
        }

        entities_hint = []
        if props.author:
            entities_hint.append(props.author)
        if props.last_modified_by:
            entities_hint.append(props.last_modified_by)

        return ParseResult(
            text_content=full_text,
            metadata=metadata,
            entities_hint=entities_hint,
            parser_name="DocxParser",
        )

    @classmethod
    def supported_mimetypes(cls) -> list[str]:
        return [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ]
